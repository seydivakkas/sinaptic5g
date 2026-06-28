from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "SINAPTIC5G_JURI_TEKNIK_RAPORU.md"
OUTPUT = ROOT / "SINAPTIC5G_JURI_TEKNIK_RAPORU.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GRAY = "555555"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF7E0"
GOLD = "7A5A00"
PALE_RED = "FDECEC"
RED = "9B1C1C"
GREEN = "1F5E3B"
WHITE = "FFFFFF"
BLACK = "000000"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=GRAY)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.10


def add_inline(paragraph, text: str, base_size=10.5, base_color=BLACK) -> None:
    token_re = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=base_size, color=base_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, color=base_color, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=max(8.2, base_size - 1), color=DARK_BLUE)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size, color=base_color)


def add_callout(doc: Document, text: str) -> None:
    fill = PALE_BLUE
    accent = DARK_BLUE
    if "UYARI" in text or "DÜRÜSTLÜK" in text:
        fill = PALE_GOLD
        accent = GOLD
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.10
    add_inline(paragraph, text, base_size=9.8, base_color=accent)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def is_numeric_column(values: list[str]) -> bool:
    nonempty = [v.strip() for v in values if v.strip()]
    if not nonempty:
        return False
    score = sum(bool(re.fullmatch(r"[\d.,%/: -]+", v.replace("ÖLÇÜLDÜ", "").strip())) for v in nonempty)
    return score >= max(1, len(nonempty) // 2)


def table_widths(rows: list[list[str]]) -> list[int]:
    cols = len(rows[0])
    if cols == 7 and rows[0][0].strip() == "Sınıf":
        inches = [1.42, 0.52, 0.66, 0.60, 0.54, 0.62, 2.14]
        return [round(v / 6.5 * TABLE_WIDTH_DXA) for v in inches[:-1]] + [TABLE_WIDTH_DXA - sum(round(v / 6.5 * TABLE_WIDTH_DXA) for v in inches[:-1])]
    if cols == 2:
        return [2700, 6660]
    max_lens = []
    for c in range(cols):
        vals = [re.sub(r"[`*]", "", row[c]) for row in rows]
        max_lens.append(max(8, min(42, max(len(v) for v in vals))))
    total = sum(max_lens)
    widths = [max(750, int(TABLE_WIDTH_DXA * value / total)) for value in max_lens]
    delta = TABLE_WIDTH_DXA - sum(widths)
    widths[-1] += delta
    if widths[-1] < 650:
        deficit = 650 - widths[-1]
        widths[-1] = 650
        widths[0] -= deficit
    return widths


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    widths = table_widths(rows)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    font_size = 8.2 if cols >= 6 else (8.8 if cols >= 4 else 9.2)
    numeric_flags = [is_numeric_column([row[c] for row in rows[1:]]) for c in range(cols)]
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (r_idx == 0 or numeric_flags[c_idx]) else WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, value, base_size=font_size, base_color=NAVY if r_idx == 0 else BLACK)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run("SİNAPTİC5G")
    set_run_font(run, size=30, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
    run = p.add_run("5G ve Yapay Zekâ ile Akıllı Yol Güvenliği Yarışması")
    set_run_font(run, size=15, color=DARK_BLUE, bold=True)
    p.add_run("\n")
    run = p.add_run("Final Teknik Raporu")
    set_run_font(run, size=20, color=BLUE, bold=True)

    meta = [
        ("Takım", "SinapticLink5G"),
        ("Takım ID", "989764"),
        ("Başvuru ID", "5205583"),
        ("Rapor tarihi", "20 Haziran 2026"),
        ("Kanıt kesimi", "20 Haziran 2026"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2200, 7160])
    for idx, (label, value) in enumerate(meta):
        set_cell_shading(table.cell(idx, 0), LIGHT_GRAY)
        p1 = table.cell(idx, 0).paragraphs[0]
        p2 = table.cell(idx, 1).paragraphs[0]
        p1.paragraph_format.space_after = Pt(1)
        p2.paragraph_format.space_after = Pt(1)
        r = p1.add_run(label)
        set_run_font(r, size=10, color=NAVY, bold=True)
        r = p2.add_run(value)
        set_run_font(r, size=10, color=BLACK)

    doc.add_paragraph().paragraph_format.space_after = Pt(16)
    add_callout(doc, "KANIT DİSİPLİNİ - ÖLÇÜLDÜ: ham kanıtla yeniden denetlenebilir. HEDEF: dış ortam veya yeni ölçüm gerektirir. UYARI: teslim öncesi kapatılması gereken tutarsızlık.")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(35)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TEKNOFEST 2026")
    set_run_font(r, size=11, color=GRAY, bold=True)
    p.add_run("\n")
    r = p.add_run("Mühendislik iddiası, kanıtı kadar güçlüdür.")
    set_run_font(r, size=10, color=GRAY, italic=True)
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph("İÇİNDEKİLER", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    entries = [
        "Yönetici Özeti",
        "1. Proje Özeti",
        "2. Veriseti Oluşturulması",
        "3. Yapay Zekâ Çözümü",
        "4. Çözümün Sınanması",
        "5. Kaynakça",
        "Ek A. Kanıtların Fiziksel Konumu",
        "Ek B. Son Mühendislik Beyanı",
    ]
    for entry in entries:
        p = doc.add_paragraph(style="List Number")
        add_inline(p, entry)
    doc.add_page_break()


def parse_markdown_body(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    start = next(i for i, line in enumerate(lines) if line.strip() == "## Yönetici Özeti")
    lines = lines[start:]
    i = 0
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = " ".join(part.strip() for part in paragraph_buffer).strip()
        p = doc.add_paragraph()
        p.paragraph_format.widow_control = True
        add_inline(p, text)
        paragraph_buffer = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            i += 1
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            flush_paragraph()
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = []
            for idx, table_line in enumerate(table_lines):
                cells = [cell.strip() for cell in table_line.strip("|").split("|")]
                if idx == 1 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                    continue
                rows.append(cells)
            if rows and len({len(row) for row in rows}) == 1:
                add_markdown_table(doc, rows)
            continue
        if stripped.startswith("> "):
            flush_paragraph()
            add_callout(doc, stripped[2:].strip())
            i += 1
            continue
        if stripped.startswith("### "):
            flush_paragraph()
            doc.add_paragraph(stripped[4:].strip(), style="Heading 3")
            i += 1
            continue
        if stripped.startswith("## "):
            flush_paragraph()
            doc.add_paragraph(stripped[3:].strip(), style="Heading 2")
            i += 1
            continue
        if stripped.startswith("# "):
            flush_paragraph()
            heading = doc.add_paragraph(stripped[2:].strip(), style="Heading 1")
            if stripped.startswith("# 2.") or stripped.startswith("# 3.") or stripped.startswith("# 4.") or stripped.startswith("# 5.") or stripped.startswith("# EK"):
                heading.paragraph_format.page_break_before = True
            i += 1
            continue
        if re.match(r"^- ", stripped):
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:].strip())
            i += 1
            continue
        number_match = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if number_match:
            flush_paragraph()
            p = doc.add_paragraph(style="List Number")
            add_inline(p, number_match.group(2))
            i += 1
            continue
        if stripped == "---":
            flush_paragraph()
            i += 1
            continue
        paragraph_buffer.append(stripped)
        i += 1
    flush_paragraph()


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.different_first_page_header_footer = True

        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("SİNAPTİC5G  |  Final Teknik Raporu")
        set_run_font(r, size=8.5, color=GRAY, bold=True)

        first_header = section.first_page_header
        first_header.paragraphs[0].text = ""

        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("Sayfa ")
        set_run_font(r, size=9, color=GRAY)
        add_page_number(p)


def build() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    doc.core_properties.title = "SİNAPTİC5G - Final Teknik Raporu"
    doc.core_properties.subject = "5G ve Yapay Zekâ ile Akıllı Yol Güvenliği Yarışması"
    doc.core_properties.author = "SinapticLink5G"
    doc.core_properties.keywords = "TEKNOFEST, 5G, YOLOv8, ONNX Runtime, WebRTC, CAMARA, QoD"
    configure_styles(doc)
    configure_sections(doc)
    add_cover(doc)
    add_toc(doc)
    parse_markdown_body(doc, markdown)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
